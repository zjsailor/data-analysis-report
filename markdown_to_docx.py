#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Markdown 转 DOCX 工具 - 增强版

支持将 Markdown 格式的数据分析报告转换为专业排版的 Word 文档。

使用方法：
    python markdown_to_docx.py input.md output.docx
"""

import sys
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

# Fix Windows encoding
if sys.platform == 'win32':
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8')


def set_font(run, cn='宋体', en='Times New Roman', size=10.5, bold=False, color=None):
    """设置中英文字体"""
    run.font.name = en
    run._element.rPr.rFonts.set(qn('w:eastAsia'), cn)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def add_text_with_formatting(para, text, cn='宋体', en='Times New Roman', size=10.5, bold=False):
    """添加带格式的文本（支持粗体和斜体）"""
    # 处理粗体 **text**
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            set_font(run, cn=cn, en=en, size=size, bold=True)
        elif part:
            run = para.add_run(part)
            set_font(run, cn=cn, en=en, size=size, bold=bold)


def create_style(doc, style_name, base_style, font_cn, font_en, size, bold=False):
    """创建自定义样式"""
    try:
        style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles[base_style] if base_style in doc.styles else None
        
        # 设置字体
        style.font.name = font_en
        style._element.rPr.rFonts.set(qn('w:eastAsia'), font_cn)
        style.font.size = Pt(size)
        style.font.bold = bold
        
        return style
    except:
        return None


def add_heading(doc, text, level):
    """添加标题"""
    # 清理标题文本（移除可能的 Markdown 标记）
    text = re.sub(r'#+\s*', '', text).strip()
    text = re.sub(r'\*\*', '', text)
    
    para = doc.add_paragraph()
    
    if level == 1:
        # 一级标题：黑体，三号，居中
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(18)
        para.paragraph_format.space_after = Pt(12)
        para.paragraph_format.line_spacing = 1.5
        run = para.add_run(text)
        set_font(run, cn='黑体', en='Arial', size=16, bold=True)
        
    elif level == 2:
        # 二级标题：黑体，小三，左对齐
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.line_spacing = 1.5
        run = para.add_run(text)
        set_font(run, cn='黑体', en='Arial', size=15, bold=True)
        
    elif level == 3:
        # 三级标题：黑体，四号，左对齐
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.line_spacing = 1.5
        run = para.add_run(text)
        set_font(run, cn='黑体', en='Arial', size=14, bold=True)
        
    elif level == 4:
        # 四级标题：黑体，小四，左对齐
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.line_spacing = 1.5
        run = para.add_run(text)
        set_font(run, cn='黑体', en='Arial', size=12, bold=True)


def add_paragraph(doc, text):
    """添加正文段落"""
    if not text.strip():
        return

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.first_line_indent = Pt(24)  # 首行缩进2字符
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(6)
    
    add_text_with_formatting(para, text, cn='宋体', en='Times New Roman', size=10.5)


def add_table(doc, lines, start_idx):
    """添加表格"""
    table_lines = []
    i = start_idx
    
    # 收集表格行
    while i < len(lines) and '|' in lines[i]:
        line = lines[i].strip()
        # 跳过分隔行
        if line and not re.match(r'^\|[\s\-:|]+\|$', line):
            table_lines.append(line)
        i += 1
    
    if not table_lines:
        return start_idx
    
    # 解析表格数据
    rows = []
    for line in table_lines:
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        if cells and any(cell for cell in cells):  # 跳过空行
            rows.append(cells)
    
    if rows:
        # 创建表格
        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        
        # 设置表格样式
        table.style = 'Table Grid'
        
        # 填充表格内容
        for row_idx, row_data in enumerate(rows):
            for col_idx, cell_text in enumerate(row_data):
                if col_idx < len(table.rows[row_idx].cells):
                    cell = table.rows[row_idx].cells[col_idx]
                    para = cell.paragraphs[0]
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # 表头加粗
                    if row_idx == 0:
                        add_text_with_formatting(para, cell_text, cn='黑体', en='Arial', size=10, bold=True)
                    else:
                        add_text_with_formatting(para, cell_text, cn='宋体', en='Times New Roman', size=9)
        
        # 添加表格前后间距
        if table.rows:
            table.rows[0].cells[0].paragraphs[0].paragraph_format.space_before = Pt(12)
    
    return i - 1


def parse_markdown(md_file, doc):
    """解析 Markdown 文件并转换为 DOCX"""
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    i = 0
    in_table = False
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        # 空行跳过
        if not line:
            i += 1
            continue
        
        # 分隔线
        if line.strip() == '---':
            i += 1
            continue
        
        # 标题（从长到短匹配）
        if line.startswith('#### '):
            add_heading(doc, line[5:], 4)
            i += 1
            
        elif line.startswith('### '):
            add_heading(doc, line[4:], 3)
            i += 1
            
        elif line.startswith('## '):
            add_heading(doc, line[3:], 2)
            i += 1
            
        elif line.startswith('# '):
            add_heading(doc, line[2:], 1)
            i += 1
        
        # 表格
        elif line.startswith('|'):
            i = add_table(doc, lines, i)
            i += 1
        
        # 列表项（无序）
        elif re.match(r'^\s*[-*]\s+', line):
            text = re.sub(r'^\s*[-*]\s+', '', line)
            para = doc.add_paragraph(style='List Bullet')
            para.paragraph_format.line_spacing = 1.5
            para.paragraph_format.left_indent = Pt(24)
            add_text_with_formatting(para, text, cn='宋体', en='Times New Roman', size=10.5)
            i += 1
        
        # 列表项（有序）
        elif re.match(r'^\s*\d+\.\s', line):
            text = re.sub(r'^\s*\d+\.\s', '', line)
            para = doc.add_paragraph(style='List Number')
            para.paragraph_format.line_spacing = 1.5
            para.paragraph_format.left_indent = Pt(24)
            add_text_with_formatting(para, text, cn='宋体', en='Times New Roman', size=10.5)
            i += 1
        
        # 正文
        else:
            add_paragraph(doc, line)
            i += 1


def setup_document(doc):
    """设置文档基本格式"""
    # 设置页面大小（A4）
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)


def main():
    if len(sys.argv) < 2:
        print("使用方法: python markdown_to_docx.py input.md [output.docx]")
        print("如果不指定输出文件名，将自动生成")
        sys.exit(1)
    
    input_file = Path(sys.argv[1])
    
    if not input_file.exists():
        print(f"错误：找不到文件 {input_file}")
        sys.exit(1)
    
    # 生成输出文件名
    if len(sys.argv) > 2:
        output_file = Path(sys.argv[2])
    else:
        output_file = input_file.with_suffix('.docx')
    
    print(f"正在转换：{input_file}")
    
    # 创建文档
    doc = Document()
    setup_document(doc)
    
    # 解析并转换
    parse_markdown(str(input_file), doc)
    
    # 保存
    doc.save(str(output_file))
    
    print(f"✅ 转换完成：{output_file}")


if __name__ == '__main__':
    main()
